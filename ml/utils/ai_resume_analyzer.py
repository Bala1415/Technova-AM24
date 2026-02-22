import os
import streamlit as st
from dotenv import load_dotenv
import requests
from google import genai
from google.genai import types
import pdfplumber
from pdf2image import convert_from_path
import pytesseract
import tempfile
import requests
import json
import math
import re


class AIResumeAnalyzer:
    def __init__(self):
        # Load environment variables
        load_dotenv()
        
        # Ollama Configuration
        self.ollama_base_url = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
        self.model_name = os.getenv("OLLAMA_MODEL", "llama3.2:3b")
        
        # Gemini Configuration (new SDK)
        self.gemini_api_key = os.getenv("GEMINI_API_KEY", "")
        try:
             self.gemini_client = genai.Client(api_key=self.gemini_api_key)
        except Exception as e:
             print(f"Failed to configure Gemini: {e}")
             self.gemini_client = None
    
    def analyze_resume_with_ollama(self, resume_text, job_description=None, job_role=None):
        """Analyze resume using Local Ollama (Llama 3.1)"""
        if not resume_text:
            return {"error": "Resume text is required for analysis."}
        
        try:
            # Construct the prompt
            system_prompt = "You are an expert resume analyst. Provide professional, detailed feedback."
            
            user_prompt = f"""
            Please analyze the following resume and provide a structured response.
            
            RESUME CONTENT:
            {resume_text}
            
            """
            
            if job_role:
                user_prompt += f"\nTARGET ROLE: {job_role}\n"
            
            if job_description:
                user_prompt += f"\nJOB DESCRIPTION:\n{job_description}\n"
                
            user_prompt += """
            
            Please provide your analysis in the following JSON format:
            {
                "resume_score": <number 0-100>,
                "ats_score": <number 0-100>,
                "summary": "<executive summary of the resume>",
                "strengths": ["<strength 1>", "<strength 2>", ...],
                "weaknesses": ["<weakness 1>", "<weakness 2>", ...],
                "missing_skills": ["<skill 1>", "<skill 2>", ...],
                "experience_analysis": "<detailed analysis of experience>",
                "education_analysis": "<analysis of education>",
                "formatting_analysis": "<analysis of formatting/ATS compatibility>",
                "recommendations": ["<recommendation 1>", "<recommendation 2>", ...]
            }
            
            Ensure the response is VALID JSON only. Do not include markdown formatting or explanations outside the JSON.
            """
            
            payload = {
                "model": self.model_name,
                "prompt": user_prompt,
                "system": system_prompt,
                "stream": False,
                "format": "json" # Llama 3 supports json mode
            }
            
            response = requests.post(f"{self.ollama_base_url}/api/generate", json=payload)
            response.raise_for_status()
            
            result = response.json()
            analysis_text = result.get("response", "")
            
            # Parse the JSON response from the LLM
            try:
                analysis_data = json.loads(analysis_text)
                
                # Normalize keys just in case
                resume_score = analysis_data.get("resume_score", 0)
                ats_score = analysis_data.get("ats_score", 0)
                
                # Format the full text analysis for the frontend/PDF report
                formatted_analysis = self._format_json_analysis_to_text(analysis_data)
                
                return {
                    "analysis": formatted_analysis,
                    "resume_score": resume_score,
                    "ats_score": ats_score,
                    "raw_json": analysis_data,
                    "model_used": self.model_name
                }
                
            except json.JSONDecodeError:
                # Fallback if valid JSON wasn't returned
                return {
                    "analysis": analysis_text,
                    "resume_score": 0, # Could use regex to extract if needed
                    "ats_score": 0,
                    "model_used": self.model_name
                }

        except Exception as e:
            return {"error": f"Ollama Analysis failed: {str(e)}"}

    def _format_json_analysis_to_text(self, data):
        """Helper to convert JSON analysis back to the readable format expected by the frontend"""
        text = ""
        text += f"## Overall Assessment\n{data.get('summary', '')}\n\n"
        
        text += "## Key Strengths\n"
        for s in data.get('strengths', []):
            text += f"- {s}\n"
        text += "\n"
        
        text += "## Areas for Improvement\n"
        for w in data.get('weaknesses', []):
            text += f"- {w}\n"
        text += "\n"
        
        text += "## Missing Skills\n"
        for s in data.get('missing_skills', []):
            text += f"- {s}\n"
        text += "\n"
        
        text += f"## Experience Analysis\n{data.get('experience_analysis', '')}\n\n"
        text += f"## Education Analysis\n{data.get('education_analysis', '')}\n\n"
        
        text += f"## Recommended Courses/Certifications\n"
        for r in data.get('recommendations', []):
             text += f"- {r}\n"
             
        text += f"\nResume Score: {data.get('resume_score', 0)}/100\n"
        text += f"ATS Score: {data.get('ats_score', 0)}/100\n"
        
        return text
    
    def extract_text_from_pdf(self, pdf_file):
        """Extract text from PDF using pdfplumber and OCR if needed"""
        text = ""
        
        # Save the uploaded file to a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
            if hasattr(pdf_file, 'getbuffer'):
                temp_file.write(pdf_file.getbuffer())
            elif hasattr(pdf_file, 'read'):
                temp_file.write(pdf_file.read())
                pdf_file.seek(0)  # Reset file pointer
            else:
                # If it's already bytes
                temp_file.write(pdf_file)
            temp_path = temp_file.name
        
        try:
            # Try direct text extraction with pdfplumber
            try:
                with pdfplumber.open(temp_path) as pdf:
                    for page in pdf.pages:
                        try:
                            # Suppress specific warnings about PDFColorSpace conversion
                            import warnings
                            with warnings.catch_warnings():
                                warnings.filterwarnings("ignore", message=".*PDFColorSpace.*")
                                warnings.filterwarnings("ignore", message=".*Cannot convert.*")
                                page_text = page.extract_text()
                                if page_text:
                                    text += page_text + "\n"
                        except Exception as e:
                            # Don't show these specific errors to the user
                            if "PDFColorSpace" not in str(e) and "Cannot convert" not in str(e):
                                print(f"Warning: Error extracting text from page with pdfplumber: {e}")
            except Exception as e:
                print(f"Warning: pdfplumber extraction failed: {e}")
                import traceback
                traceback.print_exc()

            # If pdfplumber extraction worked, return the text
            if text.strip():
                try:
                    os.unlink(temp_path)  # Clean up the temp file
                except:
                    pass
                return text.strip()
            
            # Try PyPDF2 as a fallback
            print("Info: Trying PyPDF2 extraction method...")
            try:
                import pypdf
                pdf_text = ""
                with open(temp_path, 'rb') as file:
                    pdf_reader = pypdf.PdfReader(file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            pdf_text += page_text + "\n"
                
                if pdf_text.strip():
                    try:
                        os.unlink(temp_path)  # Clean up the temp file
                    except:
                        pass
                    return pdf_text.strip()
            except Exception as e:
                print(f"Warning: PyPDF2 extraction failed: {e}")
            
            # If we got here, both extraction methods failed
            print("Warning: Standard text extraction methods failed. Your PDF might be image-based or scanned.")
            
            # Try OCR as a last resort
            try:
                # Check if we can import the required OCR libraries
                import pytesseract
                from pdf2image import convert_from_path
                
                print("Info: Attempting OCR for image-based PDF. This may take a moment...")
                
                # Check if poppler is installed
                poppler_path = None
                if os.name == 'nt':  # Windows
                    # Try to find poppler in common locations
                    possible_paths = [
                        r'C:\poppler\Library\bin',
                        r'C:\Program Files\poppler\bin',
                        r'C:\Program Files (x86)\poppler\bin',
                        r'C:\poppler\bin'
                    ]
                    for path in possible_paths:
                        if os.path.exists(path):
                            poppler_path = path
                            print(f"Info: Found Poppler at: {path}")
                            break
                    
                    if not poppler_path:
                        print("Warning: Poppler not found in common locations. Using default path: C:\\poppler\\Library\\bin")
                        poppler_path = r'C:\poppler\Library\bin'
                
                # Try to convert PDF to images
                try:
                    if poppler_path and os.name == 'nt':
                        images = convert_from_path(temp_path, poppler_path=poppler_path)
                    else:
                        images = convert_from_path(temp_path)
                    
                    # Process each image with OCR
                    ocr_text = ""
                    for i, image in enumerate(images):
                        print(f"Info: Processing page {i+1} with OCR...")
                        page_text = pytesseract.image_to_string(image)
                        ocr_text += page_text + "\n"
                    
                    if ocr_text.strip():
                        try:
                            os.unlink(temp_path)
                        except:
                            pass
                        return ocr_text.strip()
                    else:
                        print("Error: OCR extraction yielded no text. Please check if the PDF contains actual text content.")
                except Exception as e:
                    print(f"Error: PDF to image conversion failed: {e}")
                    print("Info: If you're on Windows, make sure Poppler is installed and in your PATH.")
                    print("Info: Download Poppler from: https://github.com/oschwartz10612/poppler-windows/releases/")
            except ImportError as e:
                print(f"Error: OCR libraries not available: {e}")
                print("Info: Please install the required OCR libraries: pip install pytesseract pdf2image")
            except Exception as e:
                print(f"Error: OCR processing failed: {e}")
        
        except Exception as e:
            print(f"Error: PDF processing failed: {e}")
        
        # Clean up the temp file
        try:
            os.unlink(temp_path)
        except:
            pass
        
        # If all extraction methods failed, return an empty string
        print("Error: All text extraction methods failed. Please try a different PDF or manually extract the text.")
        return ""
    
    
    def extract_text_from_docx(self, docx_file):
        """Extract text from DOCX file"""
        from docx import Document
        import io
        
        # Save the uploaded file to a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            # Handle different input types
            if isinstance(docx_file, bytes):
                temp_file.write(docx_file)
            elif hasattr(docx_file, 'getbuffer'):
                temp_file.write(docx_file.getbuffer())
            elif hasattr(docx_file, 'read'):
                temp_file.write(docx_file.read())
            else:
                temp_file.write(docx_file)
            temp_path = temp_file.name
        
        text = ""
        try:
            doc = Document(temp_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            import traceback
            traceback.print_exc()
        
        # Clean up the temp file
        try:
            os.unlink(temp_path)
        except:
            pass
            
        return text
    
    
    def analyze_resume_with_ollama(self, resume_text, job_description=None, job_role=None):
        """Analyze resume using Local Ollama (Llama 3.1)"""
        if not resume_text:
            return {"error": "Resume text is required for analysis."}
        
        try:
            # Construct the prompt
            system_prompt = "You are an expert resume analyst. Provide professional, detailed feedback."
            
            user_prompt = f"""
            Please analyze the following resume and provide a structured response.
            
            RESUME CONTENT:
            {resume_text}
            
            """
            
            if job_role:
                user_prompt += f"\nTARGET ROLE: {job_role}\n"
            
            if job_description:
                user_prompt += f"\nJOB DESCRIPTION:\n{job_description}\n"
                
            user_prompt += """
            
            Please provide your analysis in the following JSON format:
            {
                "resume_score": <number 0-100>,
                "ats_score": <number 0-100>,
                "summary": "<executive summary of the resume>",
                "strengths": ["<strength 1>", "<strength 2>", ...],
                "weaknesses": ["<weakness 1>", "<weakness 2>", ...],
                "missing_skills": ["<skill 1>", "<skill 2>", ...],
                "experience_analysis": "<detailed analysis of experience>",
                "education_analysis": "<analysis of education>",
                "formatting_analysis": "<analysis of formatting/ATS compatibility>",
                "recommendations": ["<recommendation 1>", "<recommendation 2>", ...]
            }
            
            Ensure the response is VALID JSON only. Do not include markdown formatting or explanations outside the JSON.
            """
            
            payload = {
                "model": self.model_name,
                "prompt": user_prompt,
                "system": system_prompt,
                "stream": False,
                "format": "json" # Llama 3 supports json mode
            }
            
            response = requests.post(f"{self.ollama_base_url}/api/generate", json=payload)
            response.raise_for_status()
            
            result = response.json()
            analysis_text = result.get("response", "")
            
            # Parse the JSON response from the LLM
            try:
                analysis_data = json.loads(analysis_text)
                
                # Normalize keys just in case
                resume_score = analysis_data.get("resume_score", 0)
                ats_score = analysis_data.get("ats_score", 0)
                
                # Format the full text analysis for the frontend/PDF report
                formatted_analysis = self._format_json_analysis_to_text(analysis_data)
                
                return {
                    "analysis": formatted_analysis,
                    "resume_score": resume_score,
                    "ats_score": ats_score,
                    "raw_json": analysis_data,
                    "model_used": self.model_name
                }
                
            except json.JSONDecodeError:
                # Fallback if valid JSON wasn't returned
                return {
                    "analysis": analysis_text,
                    "resume_score": 0,
                    "ats_score": 0,
                    "model_used": self.model_name
                }

        except Exception as e:
            return {"error": f"Ollama Analysis failed: {str(e)}"}

    def analyze_resume_with_gemini(self, resume_text, job_description=None, job_role=None):
        """Analyze resume using Google Gemini"""
        if not resume_text:
            return {"error": "Resume text is required for analysis."}
            
        try:
            # Construct the prompt
            system_prompt = "You are an expert resume analyst. Provide professional, detailed feedback."
            
            user_prompt = f"""
            Please analyze the following resume and provide a structured response.
            
            RESUME CONTENT:
            {resume_text}
            
            """
            
            if job_role:
                user_prompt += f"\nTARGET ROLE: {job_role}\n"
            
            if job_description:
                user_prompt += f"\nJOB DESCRIPTION:\n{job_description}\n"
                
            user_prompt += """
            
            Please provide your analysis in the following JSON format:
            {
                "resume_score": <number 0-100>,
                "ats_score": <number 0-100>,
                "summary": "<executive summary of the resume>",
                "strengths": ["<strength 1>", "<strength 2>", ...],
                "weaknesses": ["<weakness 1>", "<weakness 2>", ...],
                "missing_skills": ["<skill 1>", "<skill 2>", ...],
                "experience_analysis": "<detailed analysis of experience>",
                "education_analysis": "<analysis of education>",
                "formatting_analysis": "<analysis of formatting/ATS compatibility>",
                "recommendations": ["<recommendation 1>", "<recommendation 2>", ...]
            }
            
            Ensure the response is VALID JSON only. Do not include markdown formatting like ```json or explanations outside the JSON.
            """
            
            response = self.gemini_client.models.generate_content(
                model='gemini-2.0-flash',
                contents=user_prompt,
            )
            
            analysis_text = response.text
            
            # Remove markdown if present (Gemini sometimes adds ```json)
            if analysis_text.startswith("```json"):
                analysis_text = analysis_text[7:]
            if analysis_text.startswith("```"):
                analysis_text = analysis_text[3:]
            if analysis_text.endswith("```"):
                analysis_text = analysis_text[:-3]
            
            analysis_text = analysis_text.strip()

            # Parse the JSON response
            try:
                analysis_data = json.loads(analysis_text)
                
                # Normalize keys just in case
                resume_score = analysis_data.get("resume_score", 0)
                ats_score = analysis_data.get("ats_score", 0)
                
                # Format the full text analysis for the frontend/PDF report
                formatted_analysis = self._format_json_analysis_to_text(analysis_data)
                
                return {
                    "analysis": formatted_analysis,
                    "resume_score": resume_score,
                    "ats_score": ats_score,
                    "raw_json": analysis_data,
                    "model_used": "gemini-2.0-flash"
                }
                
            except json.JSONDecodeError:
                # Fallback if valid JSON wasn't returned
                return {
                    "analysis": analysis_text,
                    "resume_score": 0,
                    "ats_score": 0,
                    "model_used": "gemini-2.0-flash"
                }

        except Exception as e:
            return {"error": f"Gemini Analysis failed: {str(e)}"}

    def _format_json_analysis_to_text(self, data):
        """Helper to convert JSON analysis back to the readable format expected by the frontend"""
        text = ""
        text += f"## Overall Assessment\n{data.get('summary', '')}\n\n"
        
        text += "## Key Strengths\n"
        for s in data.get('strengths', []):
            text += f"- {s}\n"
        text += "\n"
        
        text += "## Areas for Improvement\n"
        for w in data.get('weaknesses', []):
            text += f"- {w}\n"
        text += "\n"
        
        text += "## Missing Skills\n"
        for s in data.get('missing_skills', []):
            text += f"- {s}\n"
        text += "\n"
        
        text += f"## Experience Analysis\n{data.get('experience_analysis', '')}\n\n"
        text += f"## Education Analysis\n{data.get('education_analysis', '')}\n\n"
        
        text += f"## Recommended Courses/Certifications\n"
        for r in data.get('recommendations', []):
             text += f"- {r}\n"
             
        text += f"\nResume Score: {data.get('resume_score', 0)}/100\n"
        text += f"ATS Score: {data.get('ats_score', 0)}/100\n"
        
        return text

    
    def generate_pdf_report(self, analysis_result, candidate_name, job_role):
        """Generate a PDF report of the analysis"""
        try:
            # Import required libraries
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.lib import colors
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, Flowable, KeepTogether
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import inch
                from reportlab.graphics.shapes import Drawing, Rect, String, Line
                from reportlab.graphics.charts.piecharts import Pie
                from reportlab.graphics.charts.barcharts import VerticalBarChart
                from reportlab.graphics.charts.linecharts import HorizontalLineChart
                from reportlab.graphics.charts.legends import Legend
                import io
                import datetime
                import math
            except ImportError as e:
                st.error(f"Error importing PDF libraries: {str(e)}")
                st.info("Please make sure reportlab is installed: pip install reportlab")
                return self.simple_generate_pdf_report(analysis_result, candidate_name, job_role)
            
            # Helper function to clean markdown formatting
            def clean_markdown(text):
                if not text:
                    return ""
                
                # Remove markdown formatting for bold and italic
                text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Remove ** for bold
                text = re.sub(r'\*(.*?)\*', r'\1', text)      # Remove * for italic
                text = re.sub(r'__(.*?)__', r'\1', text)      # Remove __ for bold
                text = re.sub(r'_(.*?)_', r'\1', text)        # Remove _ for italic
                
                # Remove markdown formatting for headers
                text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
                
                # Remove markdown formatting for links
                text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
                
                return text.strip()
            
            # Validate input data
            if not analysis_result:
                st.error("No analysis result provided for PDF generation")
                return None
                
            # Print debug info
            st.info(f"Generating PDF report for {candidate_name} targeting {job_role}")
            
            # Create a buffer for the PDF
            buffer = io.BytesIO()
            
            # Create the PDF document
            doc = SimpleDocTemplate(buffer, pagesize=letter, 
                                   leftMargin=0.5*inch, rightMargin=0.5*inch,
                                   topMargin=0.5*inch, bottomMargin=0.5*inch)
            styles = getSampleStyleSheet()
            
            # Create custom styles
            title_style = ParagraphStyle(
                'Title',
                parent=styles['Heading1'],
                fontSize=20,
                textColor=colors.darkblue,
                spaceAfter=12,
                alignment=1  # Center alignment
            )
            
            subtitle_style = ParagraphStyle(
                'Subtitle',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.darkblue,
                spaceAfter=12,
                alignment=1  # Center alignment
            )
            
            heading_style = ParagraphStyle(
                'Heading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.white,
                spaceAfter=6,
                backColor=colors.darkblue,
                borderWidth=1,
                borderColor=colors.grey,
                borderPadding=5,
                borderRadius=5,
                alignment=1  # Center alignment
            )
            
            subheading_style = ParagraphStyle(
                'SubHeading',
                parent=styles['Heading3'],
                fontSize=12,
                textColor=colors.darkblue,
                spaceAfter=6,
                borderWidth=0,
                borderPadding=0,
                borderColor=colors.grey,
                borderRadius=0
            )
            
            normal_style = ParagraphStyle(
                'Normal',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=6,
                leading=14  # Line spacing
            )
            
            list_item_style = ParagraphStyle(
                'ListItem',
                parent=normal_style,
                leftIndent=20,
                firstLineIndent=-15,
                spaceBefore=2,
                spaceAfter=2
            )
            
            # Create a gauge chart class
            class GaugeChart(Drawing):
                def __init__(self, width, height, score, max_score=100, label=""):
                    Drawing.__init__(self, width, height)
                    self.width = width
                    self.height = height
                    self._score = int(score) if score is not None else 0  # Ensure score is an integer
                    self._max_score = max_score  # Use _max_score to avoid attribute error
                    self._label = label  # Use _label instead of label to avoid attribute error
                    
                    # Determine color based on score percentage
                    score_percent = (self._score / self._max_score) * 100 if self._max_score > 0 else 0
                    if score_percent >= 80:
                        self._color = colors.green
                        self._status = "Excellent"
                    elif score_percent >= 60:
                        self._color = colors.orange
                        self._status = "Good"
                    else:
                        self._color = colors.red
                        self._status = "Needs Improvement"
                    
                    self._draw()
                
                def _draw(self):
                    # Background
                    self.add(Rect(0, 0, self.width, self.height, 
                                 fillColor=colors.white, strokeColor=None))
                    
                    # Draw gauge background (arc)
                    center_x = self.width / 2
                    center_y = self.height / 2 - 10
                    radius = min(center_x, center_y) - 10
                    
                    # Draw the gauge background
                    for i in range(0, 101, 2):
                        angle = math.radians(180 - (i * 1.8))
                        x = center_x + radius * math.cos(angle)
                        y = center_y + radius * math.sin(angle)
                        
                        # Determine color for background segments
                        if i < 60:
                            segment_color = colors.lightgrey
                        elif i < 80:
                            segment_color = colors.lightgrey
                        else:
                            segment_color = colors.lightgrey
                        
                        # Draw a small line for each segment
                        line_length = 5
                        end_x = center_x + (radius + line_length) * math.cos(angle)
                        end_y = center_y + (radius + line_length) * math.sin(angle)
                        
                        self.add(Line(x, y, end_x, end_y, strokeColor=segment_color, strokeWidth=2))
                    
                    # Draw the colored arc for the score
                    score_angle = math.radians(180 - (self._score * 1.8))
                    score_x = center_x + radius * math.cos(score_angle)
                    score_y = center_y + radius * math.sin(score_angle)
                    
                    # Draw needle
                    self.add(Line(center_x, center_y, score_x, score_y, 
                                 strokeColor=self._color, strokeWidth=3))
                    
                    # Draw center circle
                    self.add(Circle(center_x, center_y, 5, 
                                   fillColor=self._color, strokeColor=None))
                    
                    # Draw score text
                    self.add(String(center_x, center_y - 25, f"{self._score}",
                                   fontSize=20, fillColor=self._color, 
                                   textAnchor='middle', fontName='Helvetica-Bold'))
                    
                    # Draw status text
                    self.add(String(center_x, center_y - 40, self._status,
                                   fontSize=12, fillColor=colors.black, 
                                   textAnchor='middle'))
                    
                    # Draw label
                    if self._label:
                        self.add(String(center_x, self.height - 15, self._label,
                                       fontSize=12, fillColor=colors.darkblue, 
                                       textAnchor='middle', fontName='Helvetica-Bold'))
                    
                    # Draw scale markers
                    for i in range(0, 101, 20):
                        angle = math.radians(180 - (i * 1.8))
                        x = center_x + (radius - 15) * math.cos(angle)
                        y = center_y + (radius - 15) * math.sin(angle)
                        
                        self.add(String(x, y, str(i),
                                       fontSize=8, fillColor=colors.black, 
                                       textAnchor='middle'))
            
            # Create a Circle class for the gauge
            class Circle(Rect):
                def __init__(self, cx, cy, r, **kw):
                    Rect.__init__(self, cx-r, cy-r, 2*r, 2*r, **kw)
                    self.rx = self.ry = r
            
            # Create a combined gauge chart class
            class CombinedGaugeChart(Drawing):
                def __init__(self, width, height, resume_score, ats_score, max_score=100):
                    Drawing.__init__(self, width, height)
                    self.width = width
                    self.height = height
                    self._resume_score = resume_score
                    self._ats_score = ats_score
                    self._max_score = max_score
                    
                    # Calculate combined score (weighted average)
                    self._combined_score = int((self._resume_score * 0.6) + (self._ats_score * 0.4))
                    
                    # Determine color based on score percentage
                    if self._combined_score >= 80:
                        self._color = colors.green
                        self._status = "Excellent"
                    elif self._combined_score >= 60:
                        self._color = colors.orange
                        self._status = "Good"
                    else:
                        self._color = colors.red
                        self._status = "Needs Improvement"
                    
                    self._draw()
                
                def _draw(self):
                    # Background
                    self.add(Rect(0, 0, self.width, self.height, 
                                 fillColor=colors.white, strokeColor=None))
                    
                    # Draw gauge background (arc)
                    center_x = self.width / 2
                    center_y = self.height / 2
                    radius = min(center_x, center_y) - 20
                    
                    # Draw the gauge background
                    for i in range(0, 101, 2):
                        angle = math.radians(180 - (i * 1.8))
                        x = center_x + radius * math.cos(angle)
                        y = center_y + radius * math.sin(angle)
                        
                        # Determine color for background segments
                        segment_color = colors.lightgrey
                        
                        # Draw a small line for each segment
                        line_length = 5
                        end_x = center_x + (radius + line_length) * math.cos(angle)
                        end_y = center_y + (radius + line_length) * math.sin(angle)
                        
                        self.add(Line(x, y, end_x, end_y, strokeColor=segment_color, strokeWidth=2))
                    
                    # Draw the colored arc for the combined score
                    score_angle = math.radians(180 - (self._combined_score * 1.8))
                    score_x = center_x + radius * math.cos(score_angle)
                    score_y = center_y + radius * math.sin(score_angle)
                    
                    # Draw needle
                    self.add(Line(center_x, center_y, score_x, score_y, 
                                 strokeColor=self._color, strokeWidth=3))
                    
                    # Draw center circle
                    self.add(Circle(center_x, center_y, 5, 
                                   fillColor=self._color, strokeColor=None))
                    
                    # Draw combined score text
                    self.add(String(center_x, center_y - 25, f"{self._combined_score}",
                                   fontSize=24, fillColor=self._color, 
                                   textAnchor='middle', fontName='Helvetica-Bold'))
                    
                    # Draw status text
                    self.add(String(center_x, center_y - 45, self._status,
                                   fontSize=12, fillColor=colors.black, 
                                   textAnchor='middle'))
                    
                    # Draw individual scores
                    self.add(String(center_x - 60, center_y - 70, f"Resume: {self._resume_score}",
                                   fontSize=10, fillColor=colors.darkblue, 
                                   textAnchor='middle'))
                    
                    self.add(String(center_x + 60, center_y - 70, f"ATS: {self._ats_score}",
                                   fontSize=10, fillColor=colors.darkblue, 
                                   textAnchor='middle'))
                    
                    # Draw "Overall Score" label
                    self.add(String(center_x, self.height - 15, "Overall Score",
                                   fontSize=14, fillColor=colors.darkblue, 
                                   textAnchor='middle', fontName='Helvetica-Bold'))
                    
                    # Draw scale markers
                    for i in range(0, 101, 20):
                        angle = math.radians(180 - (i * 1.8))
                        x = center_x + (radius - 15) * math.cos(angle)
                        y = center_y + (radius - 15) * math.sin(angle)
                        
                        self.add(String(x, y, str(i),
                                       fontSize=8, fillColor=colors.black, 
                                       textAnchor='middle'))
            
            # Create the content
            content = []
            
            # Add a header with date
            current_date = datetime.datetime.now().strftime("%B %d, %Y")
            content.append(Paragraph(f"Resume Analysis Report", title_style))
            content.append(Paragraph(f"Generated on {current_date}", subtitle_style))
            content.append(Spacer(1, 0.25*inch))
            
            # Format candidate name - if it's just "Candidate", add a number
            if not candidate_name or candidate_name.lower() == "candidate" or candidate_name.strip() == "":
                import random
                candidate_name = f"Candidate_{random.randint(1000, 9999)}"
            
            # Add candidate name and job role in a table
            info_data = [
                ["Candidate:", candidate_name],
                ["Target Role:", job_role if job_role else "Not specified"]
            ]
            
            info_table = Table(info_data, colWidths=[1.5*inch, 5*inch])
            info_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 12),
                ('TEXTCOLOR', (0, 0), (0, -1), colors.darkblue),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
            ]))
            
            content.append(info_table)
            content.append(Spacer(1, 0.25*inch))
            
            # Analysis Content
            analysis_text = analysis_result.get("full_response", "")
            
            # Extract key sections for the executive summary
            strengths = analysis_result.get("strengths", [])
            weaknesses = analysis_result.get("weaknesses", [])
            
            # If strengths and weaknesses are not in the structured data, try to extract from text
            if not strengths:
                if "## Key Strengths" in analysis_text:
                    strengths_section = analysis_text.split("## Key Strengths")[1].split("##")[0].strip()
                    strengths = [clean_markdown(s.strip().replace("- ", "").replace("* ", "").replace("• ", "")) 
                                for s in strengths_section.split("\n") 
                                if s.strip() and (s.strip().startswith("-") or s.strip().startswith("*") or s.strip().startswith("•"))]
                
                # Try another pattern for strengths
                if not strengths and "Key Strengths" in analysis_text:
                    strengths_section = analysis_text.split("Key Strengths")[1]
                    if "Areas for Improvement" in strengths_section:
                        strengths_section = strengths_section.split("Areas for Improvement")[0]
                    
                    # Extract lines that look like list items
                    for line in strengths_section.split("\n"):
                        line = line.strip()
                        if line and (line.startswith("-") or line.startswith("*") or line.startswith("•")):
                            strengths.append(clean_markdown(line.replace("- ", "").replace("* ", "").replace("• ", "")))
                        elif line and ":" in line and not line.startswith("#"):
                            strengths.append(clean_markdown(line))

            if not weaknesses:
                if "## Areas for Improvement" in analysis_text:
                    weaknesses_section = analysis_text.split("## Areas for Improvement")[1].split("##")[0].strip()
                    weaknesses = [clean_markdown(w.strip().replace("- ", "").replace("* ", "").replace("• ", "")) 
                                 for w in weaknesses_section.split("\n") 
                                 if w.strip() and (w.strip().startswith("-") or w.strip().startswith("*") or w.strip().startswith("•"))]
                
                # Try another pattern for weaknesses
                if not weaknesses and "Areas for Improvement" in analysis_text:
                    weaknesses_section = analysis_text.split("Areas for Improvement")[1]
                    if "##" in weaknesses_section:
                        weaknesses_section = weaknesses_section.split("##")[0]
                    
                    # Extract lines that look like list items
                    for line in weaknesses_section.split("\n"):
                        line = line.strip()
                        if line and (line.startswith("-") or line.startswith("*") or line.startswith("•")):
                            weaknesses.append(clean_markdown(line.replace("- ", "").replace("* ", "").replace("• ", "")))
                        elif line and ":" in line and not line.startswith("#"):
                            weaknesses.append(clean_markdown(line))
            
            # Extract scores
            resume_score = analysis_result.get("score", 0)
            if resume_score == 0:
                # Try to get from resume_score
                resume_score = analysis_result.get("resume_score", 0)
                
                # If still 0, try to extract from the analysis text
                if resume_score == 0 and "Resume Score:" in analysis_text:
                    score_match = re.search(r'Resume Score:\s*(\d{1,3})/100', analysis_text)
                    if score_match:
                        resume_score = int(score_match.group(1))
                    else:
                        # Try another pattern
                        score_match = re.search(r'\bResume Score:\s*(\d{1,3})\b', analysis_text)
                        if score_match:
                            resume_score = int(score_match.group(1))
                        else:
                            # Try to find any number after "Resume Score:"
                            score_section = analysis_text.split("Resume Score:")[1].split("\n")[0].strip()
                            score_match = re.search(r'\b(\d{1,3})\b', score_section)
                            if score_match:
                                resume_score = int(score_match.group(1))

            # Ensure resume_score is a valid integer
            resume_score = int(resume_score) if resume_score else 0
            resume_score = max(0, min(resume_score, 100))  # Ensure it's between 0 and 100

            ats_score = analysis_result.get("ats_score", 0)
            model_used = analysis_result.get("model_used", "AI")

            # Add model used information
            model_data = [["Analysis performed by:",model_used]]
            model_table = Table(model_data, colWidths=[1.9*inch, 5*inch])
            model_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 12),
                ('TEXTCOLOR', (0, 0), (0, -1), colors.darkblue),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
            ]))

            content.append(model_table)
            content.append(Spacer(1, 0.25*inch))

            # Add score gauges
            content.append(Paragraph("Resume Evaluation", heading_style))
            content.append(Spacer(1, 0.1*inch))

            # Create a table with the gauge
            score_table_data = [
                ["Resume Score"],
                [GaugeChart(width=300, height=200, score=resume_score, max_score=100, label="Resume Score")]
            ]
            score_table = Table(score_table_data, colWidths=[6*inch])
            score_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('FONTNAME', (0, 0), (0, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (0, 0), 14),
                ('TEXTCOLOR', (0, 0), (0, 0), colors.darkblue),
                ('BOTTOMPADDING', (0, 0), (0, 0), 10),
            ]))

            content.append(score_table)
            content.append(Spacer(1, 0.25*inch))

            # Add Executive Summary section
            content.append(Paragraph("Executive Summary", heading_style))
            content.append(Spacer(1, 0.1*inch))

            # Extract overall assessment
            overall_assessment = ""
            if "## Overall Assessment" in analysis_text:
                overall_section = analysis_text.split("## Overall Assessment")[1].split("##")[0].strip()
                overall_assessment = clean_markdown(overall_section)

            content.append(Paragraph(overall_assessment, normal_style))
            content.append(Spacer(1, 0.2*inch))

            # Key Strengths and Areas for Improvement section
            content.append(Paragraph("Key Strengths and Areas for Improvement", subheading_style))
            content.append(Spacer(1, 0.1*inch))

            if strengths or weaknesses:
                # Create data for strengths and weaknesses
                sw_data = [["Key Strengths", "Areas for Improvement"]]
                
                # Get max length of strengths and weaknesses
                max_len = max(len(strengths), len(weaknesses), 1)
                
                for i in range(max_len):
                    strength = f"• {clean_markdown(strengths[i])}" if i < len(strengths) else ""
                    weakness = f"• {clean_markdown(weaknesses[i])}" if i < len(weaknesses) else ""
                    sw_data.append([
                        Paragraph(strength, list_item_style) if strength else "",
                        Paragraph(weakness, list_item_style) if weakness else ""
                    ])
                
                sw_table = Table(sw_data, colWidths=[3*inch, 3*inch])
                sw_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, 0), colors.lightgreen),
                    ('BACKGROUND', (1, 0), (1, 0), colors.salmon),
                    ('TEXTCOLOR', (0, 0), (1, 0), colors.black),
                    ('ALIGN', (0, 0), (1, 0), 'CENTER'),
                    ('FONTNAME', (0, 0), (1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (1, 0), 10),
                    ('GRID', (0, 0), (1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                
                content.append(sw_table)
            else:
                # Add empty strengths and weaknesses with a message
                empty_data = [
                    ["Key Strengths", "Areas for Improvement"],
                    [
                        Paragraph("No specific strengths identified in the analysis.", normal_style),
                        Paragraph("No specific areas for improvement identified in the analysis.", normal_style)
                    ]
                ]
                empty_table = Table(empty_data, colWidths=[3*inch, 3*inch])
                empty_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, 0), colors.lightgreen),
                    ('BACKGROUND', (1, 0), (1, 0), colors.salmon),
                    ('TEXTCOLOR', (0, 0), (1, 0), colors.black),
                    ('ALIGN', (0, 0), (1, 0), 'CENTER'),
                    ('FONTNAME', (0, 0), (1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (1, 0), 10),
                    ('GRID', (0, 0), (1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                
                content.append(empty_table)

            content.append(Spacer(1, 0.25*inch))
            
            # Add Detailed Analysis section
            content.append(Paragraph("Detailed Analysis", heading_style))
            content.append(Spacer(1, 0.1*inch))
            
            # Parse the markdown-like content
            sections = analysis_text.split("##")
            
            # Define sections to include in detailed analysis
            detailed_sections = [
                "Professional Profile Analysis",
                "Skills Analysis",
                "Experience Analysis",
                "Education Analysis",
                "ATS Optimization Assessment",
                "Role Alignment Analysis",
                "Job Match Analysis"
            ]
            
            for section in sections:
                if not section.strip():
                    continue
                
                # Extract section title and content
                lines = section.strip().split("\n")
                section_title = lines[0].strip()
                
                # Skip sections we don't want in the detailed analysis
                if section_title not in detailed_sections and section_title != "Overall Assessment":
                    continue
                
                # Skip Overall Assessment as we've already included it
                if section_title == "Overall Assessment":
                    continue
                
                section_content = "\n".join(lines[1:]).strip()
                
                # Add section title
                content.append(Paragraph(section_title, subheading_style))
                content.append(Spacer(1, 0.1*inch))
                
                # Process content based on section
                if section_title == "Skills Analysis":
                    # Extract current and missing skills
                    current_skills = []
                    missing_skills = []
                    
                    if "Current Skills" in section_content:
                        current_part = section_content.split("Current Skills")[1]
                        if "Missing Skills" in current_part:
                            current_part = current_part.split("Missing Skills")[0]
                        
                        for line in current_part.split("\n"):
                            if line.strip() and ("-" in line or "*" in line or "•" in line):
                                skill = line.replace("-", "").replace("*", "").replace("•", "").strip()
                                if skill:
                                    current_skills.append(skill)
                    
                    if "Missing Skills" in section_content:
                        missing_part = section_content.split("Missing Skills")[1]
                        for line in missing_part.split("\n"):
                            if line.strip() and ("-" in line or "*" in line or "•" in line):
                                skill = line.replace("-", "").replace("*", "").replace("•", "").strip()
                                if skill:
                                    missing_skills.append(skill)
                    
                    # Create skills table with better formatting
                    if current_skills or missing_skills:
                        # Create paragraphs for each skill to ensure proper wrapping
                        current_skill_paragraphs = [Paragraph(skill, normal_style) for skill in current_skills]
                        missing_skill_paragraphs = [Paragraph(skill, normal_style) for skill in missing_skills]
                        
                        # Make sure both lists have the same length
                        max_len = max(len(current_skill_paragraphs), len(missing_skill_paragraphs))
                        current_skill_paragraphs.extend([Paragraph("", normal_style)] * (max_len - len(current_skill_paragraphs)))
                        missing_skill_paragraphs.extend([Paragraph("", normal_style)] * (max_len - len(missing_skill_paragraphs)))
                        
                        # Create data for the table
                        data = [["Current Skills", "Missing Skills"]]
                        for i in range(max_len):
                            data.append([current_skill_paragraphs[i], missing_skill_paragraphs[i]])
                        
                        # Create the table with fixed column widths
                        table = Table(data, colWidths=[3*inch, 3*inch])
                        table.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (1, 0), colors.lightgreen),
                            ('TEXTCOLOR', (0, 0), (1, 0), colors.black),
                            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                            ('GRID', (0, 0), (-1, -1), 1, colors.black),
                            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                            ('LEFTPADDING', (0, 0), (-1, -1), 10),
                            ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                        ]))
                        
                        content.append(table)
                    
                    # We no longer need to add skill proficiency outside the table
                    # as it's now included in the table itself
                elif section_title == "ATS Optimization Assessment":
                    # Special handling for ATS Optimization Assessment
                    ats_score_line = ""
                    ats_content = []
                    
                    # Extract ATS score if present
                    for line in section_content.split("\n"):
                        if "ATS Score:" in line:
                            ats_score_line = clean_markdown(line)
                        elif line.strip():
                            # Check if it's a list item
                            if line.strip().startswith("-") or line.strip().startswith("*") or line.strip().startswith("•"):
                                ats_content.append("• " + clean_markdown(line.strip()[1:].strip()))
                            else:
                                ats_content.append(clean_markdown(line))
                    
                    # Add ATS score line if found
                    if ats_score_line:
                        content.append(Paragraph(ats_score_line, normal_style))
                        content.append(Spacer(1, 0.1*inch))
                    
                    # Add the rest of the ATS content
                    for para in ats_content:
                        if para.startswith("• "):
                            content.append(Paragraph(para, list_item_style))
                        else:
                            content.append(Paragraph(para, normal_style))
                else:
                    # Process regular paragraphs
                    paragraphs = section_content.split("\n")
                    for para in paragraphs:
                        if para.strip():
                            # Check if it's a list item
                            if para.strip().startswith("-") or para.strip().startswith("*") or para.strip().startswith("•"):
                                para = "• " + clean_markdown(para.strip()[1:].strip())
                                content.append(Paragraph(para, list_item_style))
                            else:
                                content.append(Paragraph(clean_markdown(para), normal_style))
                
                content.append(Spacer(1, 0.2*inch))
            
            # Add course recommendations
            course_recommendations = []
            
            # Try to get course recommendations from different sources
            if "suggestions" in analysis_result:
                course_recommendations = analysis_result.get("suggestions", [])
            
            # If still no recommendations, try to extract from text
            if not course_recommendations and "## Recommended Courses" in analysis_text:
                recommendations_section = analysis_text.split("## Recommended Courses")[1].split("##")[0].strip()
                course_recommendations = [clean_markdown(r.strip().replace("- ", "").replace("* ", "").replace("• ", "")) 
                              for r in recommendations_section.split("\n") 
                              if r.strip() and (r.strip().startswith("-") or r.strip().startswith("*") or r.strip().startswith("•"))]
            
            # Try another pattern for course recommendations
            if not course_recommendations and "Recommended Courses" in analysis_text:
                recommendations_section = analysis_text.split("Recommended Courses")[1]
                if "##" in recommendations_section:
                    recommendations_section = recommendations_section.split("##")[0]
                
                # Extract lines that look like list items
                for line in recommendations_section.split("\n"):
                    line = line.strip()
                    if line and ":" in line and not line.startswith("#"):
                        course_recommendations.append(clean_markdown(line))
            
            content.append(Paragraph("Recommended Courses & Certifications", subheading_style))
            
            if course_recommendations:
                # Create a table for course recommendations with better formatting
                course_data = [["Recommended Courses & Certifications"]]  # Add header row
                
                for course in course_recommendations:
                    # Clean the course text and ensure it doesn't have any markdown formatting
                    cleaned_course = clean_markdown(course)
                    course_data.append([Paragraph(f"• {cleaned_course}", list_item_style)])
                
                course_table = Table(course_data, colWidths=[6*inch])
                course_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, 0), colors.lightblue),
                    ('TEXTCOLOR', (0, 0), (0, 0), colors.black),
                    ('ALIGN', (0, 0), (0, 0), 'CENTER'),  # Center the header
                    ('ALIGN', (0, 1), (0, -1), 'LEFT'),   # Left-align the content
                    ('FONTNAME', (0, 0), (0, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (0, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (0, 0), 10),
                    ('GRID', (0, 0), (0, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (0, -1), 'TOP'),
                ]))
                
                content.append(course_table)
            else:
                # If still no recommendations, add a text section instead of generic courses
                content.append(Paragraph("Based on your resume and target role, consider the following types of courses and certifications:", normal_style))
                content.append(Spacer(1, 0.1*inch))
                
                # Add role-specific recommendations based on job_role
                role_specific_courses = []
                if "data" in job_role.lower() or "scientist" in job_role.lower() or "analyst" in job_role.lower():
                    role_specific_courses = [
                        "Data Science Specialization (Coursera/edX)",
                        "Machine Learning (Coursera/edX)",
                        "Deep Learning Specialization (Coursera)",
                        "Big Data Technologies (Cloud Provider Certifications)",
                        "Statistical Modeling and Inference",
                        "Data Visualization with Tableau/Power BI"
                    ]
                elif "developer" in job_role.lower() or "engineer" in job_role.lower() or "programming" in job_role.lower():
                    role_specific_courses = [
                        "Full Stack Web Development (Udemy/Coursera)",
                        "Cloud Certifications (AWS/Azure/GCP)",
                        "DevOps and CI/CD Pipelines",
                        "Software Architecture and Design Patterns",
                        "Agile and Scrum Methodologies",
                        "Mobile App Development"
                    ]
                elif "security" in job_role.lower() or "cyber" in job_role.lower():
                    role_specific_courses = [
                        "Certified Information Systems Security Professional (CISSP)",
                        "Certified Ethical Hacker (CEH)",
                        "CompTIA Security+",
                        "Offensive Security Certified Professional (OSCP)",
                        "Cloud Security Certifications",
                        "Security Operations and Incident Response"
                    ]
                else:
                    # Generic professional development courses
                    role_specific_courses = [
                        "LinkedIn Learning - Professional Skills Development",
                        "Coursera - Career Development Specialization",
                        "Udemy - Job Interview Skills Training",
                        "Project Management Professional (PMP)",
                        "Leadership and Management Skills",
                        "Technical Writing and Communication"
                    ]
                
                # Create a table for role-specific courses
                course_data = []
                for course in role_specific_courses:
                    course_data.append([Paragraph(f"• {clean_markdown(course)}", list_item_style)])
                
                course_table = Table(course_data, colWidths=[6*inch])
                course_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, 0), colors.lightblue),
                    ('TEXTCOLOR', (0, 0), (0, 0), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica'),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ]))
                
                content.append(course_table)
            
            content.append(Spacer(1, 0.2*inch))
            
            # Add footer with page numbers
            def add_page_number(canvas, doc):
                canvas.saveState()
                canvas.setFont('Helvetica', 9)
                page_num = canvas.getPageNumber()
                text = f"Page {page_num}"
                canvas.drawRightString(7.5*inch, 0.25*inch, text)
                
                # Add generation date at the bottom
                canvas.setFont('Helvetica', 9)
                date_text = f"Generated on: {datetime.datetime.now().strftime('%B %d, %Y')}"
                canvas.drawString(0.5*inch, 0.25*inch, date_text)
                
                canvas.restoreState()
            
            # Build the PDF
            doc.build(content, onFirstPage=add_page_number, onLaterPages=add_page_number)
            
            # Get the PDF from the buffer
            buffer.seek(0)
            return buffer
        
        except Exception as e:
            st.error(f"Error generating simple PDF report: {str(e)}")
            import traceback
            st.code(traceback.format_exc())
            return None
            
    def extract_skills_from_analysis(self, analysis_text):
        """Extract skills from the analysis text"""
        skills = []
        
        try:
            if "Current Skills" in analysis_text:
                skills_section = analysis_text.split("Current Skills")[1]
                if "##" in skills_section:
                    skills_section = skills_section.split("##")[0]
                
                for line in skills_section.split("\n"):
                    if line.strip() and ("-" in line or "*" in line or "•" in line):
                        skill = line.replace("-", "").replace("*", "").replace("•", "").strip()
                        if skill:
                            skills.append(skill)
        except Exception as e:
            st.warning(f"Error extracting skills: {str(e)}")
        
        return skills
        
    def extract_missing_skills_from_analysis(self, analysis_text):
        """Extract missing skills from the analysis text"""
        missing_skills = []
        
        try:
            if "Missing Skills" in analysis_text:
                missing_section = analysis_text.split("Missing Skills")[1]
                if "##" in missing_section:
                    missing_section = missing_section.split("##")[0]
                
                for line in missing_section.split("\n"):
                    if line.strip() and ("-" in line or "*" in line or "•" in line):
                        skill = line.replace("-", "").replace("*", "").replace("•", "").strip()
                        if skill:
                            missing_skills.append(skill)
        except Exception as e:
            st.warning(f"Error extracting missing skills: {str(e)}")
        
        return missing_skills
    
    def _extract_score_from_text(self, analysis_text):
        """Extract the resume score from the analysis text"""
        try:
            # Look for the Resume Score section
            if "## Resume Score" in analysis_text:
                score_section = analysis_text.split("## Resume Score")[1].strip()
                # Extract the first number found
                score_match = re.search(r'Resume Score:\s*(\d{1,3})/100', score_section)
                if score_match:
                    score = int(score_match.group(1))
                    # Ensure score is within valid range
                    return max(0, min(score, 100))
                
                # Try another pattern if the first one doesn't match
                score_match = re.search(r'\b(\d{1,3})\b', score_section)
                if score_match:
                    score = int(score_match.group(1))
                    # Ensure score is within valid range
                    return max(0, min(score, 100))
            
            # If no score found in Resume Score section, try to find it elsewhere
            score_match = re.search(r'Resume Score:\s*(\d{1,3})/100', analysis_text)
            if score_match:
                score = int(score_match.group(1))
                return max(0, min(score, 100))
                
            return 0
        except Exception as e:
            print(f"Error extracting score: {str(e)}")
            return 0
            
    def _extract_ats_score_from_text(self, analysis_text):
        """Extract the ATS score from the analysis text"""
        try:
            # Look for the ATS Score in the ATS Optimization Assessment section
            if "## ATS Optimization Assessment" in analysis_text:
                ats_section = analysis_text.split("## ATS Optimization Assessment")[1].split("##")[0].strip()
                # Extract the score using regex
                score_match = re.search(r'ATS Score:\s*(\d{1,3})/100', ats_section)
                if score_match:
                    score = int(score_match.group(1))
                    # Ensure score is within valid range
                    return max(0, min(score, 100))
            return 0
        except Exception as e:
            print(f"Error extracting ATS score: {str(e)}")
            return 0
            
    def analyze_resume(self, resume_text, job_role=None, role_info=None, model="Google Gemini"):
        """
        Analyze a resume using the specified AI model
        
        Parameters:
        - resume_text: The text content of the resume
        - job_role: The target job role
        - role_info: Additional information about the job role
        - model: The AI model to use ("Google Gemini" or "Anthropic Claude")
        
        Returns:
        - Dictionary containing analysis results
        """
        import traceback
        
        try:
            job_description = None
            if role_info:
                job_description = f"""
                Role: {job_role}
                Description: {role_info.get('description', '')}
                Required Skills: {', '.join(role_info.get('required_skills', []))}
                """
            
            # Choose the appropriate model for analysis
            if model == "Google Gemini":
                result = self.analyze_resume_with_gemini(resume_text, job_description, job_role)
                model_used = "Google Gemini"
            elif model == "Anthropic Claude":
                result = self.analyze_resume_with_anthropic(resume_text, job_description, job_role)
                # Get the actual model used from the result
                model_used = result.get("model_used", "Anthropic Claude")
            else:
                # Default to Gemini if model not recognized
                result = self.analyze_resume_with_gemini(resume_text, job_description, job_role)
                model_used = "Google Gemini"
            
            # Process the result to extract structured information
            analysis_text = result.get("analysis", "")
            
            # Extract strengths
            strengths = []
            if "## Key Strengths" in analysis_text:
                strengths_section = analysis_text.split("## Key Strengths")[1].split("##")[0].strip()
                strengths = [clean_markdown(s.strip().replace("- ", "").replace("* ", "").replace("• ", "")) 
                            for s in strengths_section.split("\n") 
                            if s.strip() and (s.strip().startswith("-") or s.strip().startswith("*") or s.strip().startswith("•"))]
            
            # Extract weaknesses/areas for improvement
            weaknesses = []
            if "## Areas for Improvement" in analysis_text:
                weaknesses_section = analysis_text.split("## Areas for Improvement")[1].split("##")[0].strip()
                weaknesses = [clean_markdown(w.strip().replace("- ", "").replace("* ", "").replace("• ", "")) 
                             for w in weaknesses_section.split("\n") 
                             if w.strip() and (w.strip().startswith("-") or w.strip().startswith("*") or w.strip().startswith("•"))]
            
            # Extract suggestions/recommendations
            suggestions = []
            if "## Recommended Courses" in analysis_text:
                suggestions_section = analysis_text.split("## Recommended Courses")[1].split("##")[0].strip()
                suggestions = [clean_markdown(s.strip().replace("- ", "").replace("* ", "").replace("• ", "")) 
                                 for s in suggestions_section.split("\n") 
                                 if s.strip() and (s.strip().startswith("-") or s.strip().startswith("*") or s.strip().startswith("•"))]
            
            # Extract score
            score = result.get("resume_score", 0)
            if not score:
                score = self._extract_score_from_text(analysis_text)
            
            # Extract ATS score
            ats_score = self._extract_ats_score_from_text(analysis_text)
            
            # Return structured analysis
            return {
                "score": score,
                "ats_score": ats_score,
                "strengths": strengths,
                "weaknesses": weaknesses,
                "suggestions": suggestions,
                "full_response": analysis_text,
                "model_used": model_used
            }
            
        except Exception as e:
            print(f"Error in analyze_resume: {str(e)}")
            print(traceback.format_exc())
            return {
                "error": f"Analysis failed: {str(e)}",
                "score": 0,
                "ats_score": 0,
                "strengths": ["Unable to analyze resume due to an error."],
                "weaknesses": ["Unable to analyze resume due to an error."],
                "suggestions": ["Try again with a different model or check your resume format."],
                "full_response": f"Error: {str(e)}",
                "model_used": "Error"
            } 

    def simple_generate_pdf_report(self, analysis_result, candidate_name, job_role):
        """Generate a simple PDF report without complex charts as a fallback"""
        try:
            # Import required libraries
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.lib import colors
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, Flowable, KeepTogether
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import inch
                from reportlab.graphics.shapes import Drawing, Rect, String, Line
                from reportlab.graphics.charts.piecharts import Pie
                from reportlab.graphics.charts.barcharts import VerticalBarChart
                from reportlab.graphics.charts.linecharts import HorizontalLineChart
                from reportlab.graphics.charts.legends import Legend
                import io
                import datetime
                import math
            except ImportError as e:
                st.error(f"Error importing PDF libraries: {str(e)}")
                st.info("Please make sure reportlab is installed: pip install reportlab")
                return None
            
            # Helper function to clean markdown formatting
            def clean_markdown(text):
                if not text:
                    return ""
                
                # Remove markdown formatting for bold and italic
                text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Remove ** for bold
                text = re.sub(r'\*(.*?)\*', r'\1', text)      # Remove * for italic
                text = re.sub(r'__(.*?)__', r'\1', text)      # Remove __ for bold
                text = re.sub(r'_(.*?)_', r'\1', text)        # Remove _ for italic
                
                # Remove markdown formatting for headers
                text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
                
                # Remove markdown formatting for links
                text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
                
                return text.strip()
            
            # Validate input data
            if not analysis_result:
                st.error("No analysis result provided for PDF generation")
                return None
                
            # Create a buffer for the PDF
            buffer = io.BytesIO()
            
            # Create the PDF document
            doc = SimpleDocTemplate(buffer, pagesize=letter, 
                                   leftMargin=0.5*inch, rightMargin=0.5*inch,
                                   topMargin=0.5*inch, bottomMargin=0.5*inch)
            styles = getSampleStyleSheet()
            
            # Create custom styles
            title_style = ParagraphStyle(
                'Title',
                parent=styles['Heading1'],
                fontSize=20,
                textColor=colors.darkblue,
                spaceAfter=12,
                alignment=1  # Center alignment
            )
            
            subtitle_style = ParagraphStyle(
                'Subtitle',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.darkblue,
                spaceAfter=12,
                alignment=1  # Center alignment
            )
            
            heading_style = ParagraphStyle(
                'Heading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.white,
                spaceAfter=6,
                backColor=colors.darkblue,
                borderWidth=1,
                borderColor=colors.grey,
                borderPadding=5,
                borderRadius=5,
                alignment=1  # Center alignment
            )
            
            subheading_style = ParagraphStyle(
                'SubHeading',
                parent=styles['Heading3'],
                fontSize=12,
                textColor=colors.darkblue,
                spaceAfter=6
            )
            
            normal_style = ParagraphStyle(
                'Normal',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=6,
                leading=14  # Line spacing
            )
            
            list_item_style = ParagraphStyle(
                'ListItem',
                parent=normal_style,
                leftIndent=20,
                firstLineIndent=-15,
                spaceBefore=2,
                spaceAfter=2
            )
            
            # Create a simple gauge chart class
            class SimpleGaugeChart(Flowable):
                def __init__(self, score, width=300, height=200, label="Resume Score"):
                    Flowable.__init__(self)
                    self.score = int(score) if score is not None else 0  # Ensure score is an integer
                    self.width = width
                    self.height = height
                    self.label = label
                    
                    # Determine color based on score percentage
                    if self.score >= 80:
                        self.color = colors.green
                        self.status = "Excellent"
                    elif self.score >= 60:
                        self.color = colors.orange
                        self.status = "Good"
                    else:
                        self.color = colors.red
                        self.status = "Needs Improvement"
                
                def draw(self):
                    # Draw the gauge
                    canvas = self.canv
                    canvas.saveState()
                    
                    # Draw gauge background (semi-circle)
                    center_x = self.width / 2
                    center_y = self.height / 2
                    radius = min(center_x, center_y) - 30
                    
                    # Draw the gauge background
                    canvas.setFillColor(colors.lightgrey)
                    canvas.setStrokeColor(colors.grey)
                    canvas.setLineWidth(1)
                    
                    # Draw the semi-circle background
                    p = canvas.beginPath()
                    p.moveTo(center_x, center_y)
                    p.arcTo(center_x - radius, center_y - radius, center_x + radius, center_y + radius, 0, 180)
                    p.lineTo(center_x, center_y)
                    p.close()
                    canvas.drawPath(p, fill=1, stroke=1)
                    
                    # Draw the colored arc for the score
                    if self.score > 0:  # Only draw if score > 0
                        angle = 180 * self.score / 100
                        p = canvas.beginPath()
                        p.moveTo(center_x, center_y)
                        p.arcTo(center_x - radius, center_y - radius, center_x + radius, center_y + radius, 180, 180-angle)
                        p.lineTo(center_x, center_y)
                        p.close()
                        canvas.setFillColor(self.color)
                        canvas.drawPath(p, fill=1, stroke=0)
                    
                    # Draw score text
                    canvas.setFillColor(self.color)
                    canvas.setFont("Helvetica-Bold", 24)
                    canvas.drawCentredString(center_x, center_y - 15, f"{self.score}")
                    
                    # Draw status text
                    canvas.setFillColor(self.color)
                    canvas.setFont("Helvetica", 12)
                    canvas.drawCentredString(center_x, center_y - 35, self.status)
                    
                    # Draw "Resume Score" label
                    canvas.setFillColor(colors.darkblue)
                    canvas.setFont("Helvetica-Bold", 14)
                    canvas.drawCentredString(center_x, self.height - 20, self.label)
                    
                    # Draw scale markers
                    canvas.setStrokeColor(colors.black)
                    canvas.setLineWidth(1)
                    for i in range(0, 101, 20):
                        angle_rad = math.radians(180 - (i * 1.8))
                        x = center_x + radius * math.cos(angle_rad)
                        y = center_y + radius * math.sin(angle_rad)
                        
                        # Draw tick marks
                        x2 = center_x + (radius - 5) * math.cos(angle_rad)
                        y2 = center_y + (radius - 5) * math.sin(angle_rad)
                        canvas.line(x, y, x2, y2)
                        
                        # Draw numbers
                        canvas.setFont("Helvetica", 8)
                        num_x = center_x + (radius - 15) * math.cos(angle_rad)
                        num_y = center_y + (radius - 15) * math.sin(angle_rad)
                        canvas.drawCentredString(num_x, num_y, str(i))
                    
                    canvas.restoreState()
                
                def wrap(self, availWidth, availHeight):
                    return (self.width, self.height)
            
            # Create the content
            content = []
            
            # Add a header with date
            current_date = datetime.datetime.now().strftime("%B %d, %Y")
            content.append(Paragraph(f"Resume Analysis Report", title_style))
            content.append(Paragraph(f"Generated on {current_date}", subtitle_style))
            content.append(Spacer(1, 0.25*inch))
            
            # Format candidate name - if it's just "Candidate", add a number
            if not candidate_name or candidate_name.lower() == "candidate" or candidate_name.strip() == "":
                import random
                candidate_name = f"Candidate_{random.randint(1000, 9999)}"
            
            # Add candidate name and job role in a table
            info_data = [
                ["Candidate:", candidate_name],
                ["Target Role:", job_role if job_role else "Not specified"]
            ]
            
            info_table = Table(info_data, colWidths=[1.5*inch, 5*inch])
            info_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 12),
                ('TEXTCOLOR', (0, 0), (0, -1), colors.darkblue),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
            ]))
            
            content.append(info_table)
            content.append(Spacer(1, 0.25*inch))
            
            # Add model used information with proper spacing
            model_used = analysis_result.get("model_used", "AI")
            model_data = [["Analysis performed by:\u2003\u2003\u2003", "", model_used]]
            model_table = Table(model_data, colWidths=[3.5*inch, 1*inch, 5*inch])
            model_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 12),
                ('TEXTCOLOR', (0, 0), (0, -1), colors.darkblue),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 20),
            ]))
            
            content.append(model_table)
            content.append(Spacer(1, 0.25*inch))
            
            # Add Resume Evaluation section
            content.append(Paragraph("Resume Evaluation", heading_style))
            content.append(Spacer(1, 0.1*inch))
            
            # Extract scores
            resume_score = analysis_result.get("score", 0)
            if resume_score == 0:
                # Try to get from resume_score
                resume_score = analysis_result.get("resume_score", 0)
                
                # If still 0, try to extract from the analysis text
                if resume_score == 0 and "Resume Score:" in analysis_text:
                    score_match = re.search(r'Resume Score:\s*(\d{1,3})/100', analysis_text)
                    if score_match:
                        resume_score = int(score_match.group(1))
                    else:
                        # Try another pattern
                        score_match = re.search(r'\bResume Score:\s*(\d{1,3})\b', analysis_text)
                        if score_match:
                            resume_score = int(score_match.group(1))
                        else:
                            # Try to find any number after "Resume Score:"
                            score_section = analysis_text.split("Resume Score:")[1].split("\n")[0].strip()
                            score_match = re.search(r'\b(\d{1,3})\b', score_section)
                            if score_match:
                                resume_score = int(score_match.group(1))

            # Ensure resume_score is a valid integer
            resume_score = int(resume_score) if resume_score else 0
            resume_score = max(0, min(resume_score, 100))  # Ensure it's between 0 and 100

            # Create a table with the simple gauge
            score_table_data = [
                ["Resume Score"],
                [SimpleGaugeChart(score=resume_score, width=300, height=200, label="Resume Score")]
            ]
            
            score_table = Table(score_table_data, colWidths=[6*inch])
            score_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('FONTNAME', (0, 0), (0, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (0, 0), 14),
                ('TEXTCOLOR', (0, 0), (0, 0), colors.darkblue),
                ('BOTTOMPADDING', (0, 0), (0, 0), 10),
            ]))
            
            content.append(score_table)
            content.append(Spacer(1, 0.25*inch))
            
            # Add Executive Summary section
            content.append(Paragraph("Executive Summary", heading_style))
            content.append(Spacer(1, 0.1*inch))
            
            # Extract overall assessment
            analysis_text = analysis_result.get("full_response", "")
            if not analysis_text:
                analysis_text = analysis_result.get("analysis", "")
                
            overall_assessment = ""
            if "## Overall Assessment" in analysis_text:
                overall_section = analysis_text.split("## Overall Assessment")[1].split("##")[0].strip()
                overall_assessment = clean_markdown(overall_section)
            
            content.append(Paragraph(overall_assessment, normal_style))
            content.append(Spacer(1, 0.2*inch))
            
            # Key Strengths and Areas for Improvement section
            content.append(Paragraph("Key Strengths and Areas for Improvement", subheading_style))
            content.append(Spacer(1, 0.1*inch))

            if strengths or weaknesses:
                # Create data for strengths and weaknesses
                sw_data = [["Key Strengths", "Areas for Improvement"]]
                
                # Get max length of strengths and weaknesses
                max_len = max(len(strengths), len(weaknesses), 1)
                
                for i in range(max_len):
                    strength = f"• {clean_markdown(strengths[i])}" if i < len(strengths) else ""
                    weakness = f"• {clean_markdown(weaknesses[i])}" if i < len(weaknesses) else ""
                    sw_data.append([
                        Paragraph(strength, list_item_style) if strength else "",
                        Paragraph(weakness, list_item_style) if weakness else ""
                    ])
                
                sw_table = Table(sw_data, colWidths=[3*inch, 3*inch])
                sw_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, 0), colors.lightgreen),
                    ('BACKGROUND', (1, 0), (1, 0), colors.salmon),
                    ('TEXTCOLOR', (0, 0), (1, 0), colors.black),
                    ('ALIGN', (0, 0), (1, 0), 'CENTER'),
                    ('FONTNAME', (0, 0), (1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (1, 0), 10),
                    ('GRID', (0, 0), (1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                
                content.append(sw_table)
            else:
                # Add empty strengths and weaknesses with a message
                empty_data = [
                    ["Key Strengths", "Areas for Improvement"],
                    [
                        Paragraph("No specific strengths identified in the analysis.", normal_style),
                        Paragraph("No specific areas for improvement identified in the analysis.", normal_style)
                    ]
                ]
                empty_table = Table(empty_data, colWidths=[3*inch, 3*inch])
                empty_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, 0), colors.lightgreen),
                    ('BACKGROUND', (1, 0), (1, 0), colors.salmon),
                    ('TEXTCOLOR', (0, 0), (1, 0), colors.black),
                    ('ALIGN', (0, 0), (1, 0), 'CENTER'),
                    ('FONTNAME', (0, 0), (1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (1, 0), 10),
                    ('GRID', (0, 0), (1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                
                content.append(empty_table)

            content.append(Spacer(1, 0.25*inch))
            
            # Use the process_sections method to handle detailed analysis
            content = self.process_sections(analysis_text, content, normal_style, list_item_style, subheading_style, heading_style, clean_markdown)
            
            # Add course recommendations
            course_recommendations = []
            
            # Try to get course recommendations from different sources
            if "suggestions" in analysis_result:
                course_recommendations = analysis_result.get("suggestions", [])
            
            # If still no recommendations, try to extract from text
            if not course_recommendations and "## Recommended Courses" in analysis_text:
                recommendations_section = analysis_text.split("## Recommended Courses")[1].split("##")[0].strip()
                course_recommendations = [clean_markdown(r.strip().replace("- ", "").replace("* ", "").replace("• ", "")) 
                              for r in recommendations_section.split("\n") 
                              if r.strip() and (r.strip().startswith("-") or r.strip().startswith("*") or r.strip().startswith("•"))]
            
            # Try another pattern for course recommendations
            if not course_recommendations and "Recommended Courses" in analysis_text:
                recommendations_section = analysis_text.split("Recommended Courses")[1]
                if "##" in recommendations_section:
                    recommendations_section = recommendations_section.split("##")[0]
                
                # Extract lines that look like list items
                for line in recommendations_section.split("\n"):
                    line = line.strip()
                    if line and ":" in line and not line.startswith("#"):
                        course_recommendations.append(clean_markdown(line))
            
            content.append(Paragraph("Recommended Courses & Certifications", subheading_style))
            
            if course_recommendations:
                # Create a table for course recommendations with better formatting
                course_data = [["Recommended Courses & Certifications"]]  # Add header row
                
                for course in course_recommendations:
                    # Clean the course text and ensure it doesn't have any markdown formatting
                    cleaned_course = clean_markdown(course)
                    course_data.append([Paragraph(f"• {cleaned_course}", list_item_style)])
                
                course_table = Table(course_data, colWidths=[6*inch])
                course_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, 0), colors.lightblue),
                    ('TEXTCOLOR', (0, 0), (0, 0), colors.black),
                    ('ALIGN', (0, 0), (0, 0), 'CENTER'),  # Center the header
                    ('ALIGN', (0, 1), (0, -1), 'LEFT'),   # Left-align the content
                    ('FONTNAME', (0, 0), (0, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (0, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (0, 0), 10),
                    ('GRID', (0, 0), (0, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (0, -1), 'TOP'),
                ]))
                
                content.append(course_table)
            else:
                # If still no recommendations, add a text section instead of generic courses
                content.append(Paragraph("Based on your resume and target role, consider the following types of courses and certifications:", normal_style))
                content.append(Spacer(1, 0.1*inch))
                
                # Add role-specific recommendations based on job_role
                role_specific_courses = []
                if "data" in job_role.lower() or "scientist" in job_role.lower() or "analyst" in job_role.lower():
                    role_specific_courses = [
                        "Data Science Specialization (Coursera/edX)",
                        "Machine Learning (Coursera/edX)",
                        "Deep Learning Specialization (Coursera)",
                        "Big Data Technologies (Cloud Provider Certifications)",
                        "Statistical Modeling and Inference",
                        "Data Visualization with Tableau/Power BI"
                    ]
                elif "developer" in job_role.lower() or "engineer" in job_role.lower() or "programming" in job_role.lower():
                    role_specific_courses = [
                        "Full Stack Web Development (Udemy/Coursera)",
                        "Cloud Certifications (AWS/Azure/GCP)",
                        "DevOps and CI/CD Pipelines",
                        "Software Architecture and Design Patterns",
                        "Agile and Scrum Methodologies",
                        "Mobile App Development"
                    ]
                elif "security" in job_role.lower() or "cyber" in job_role.lower():
                    role_specific_courses = [
                        "Certified Information Systems Security Professional (CISSP)",
                        "Certified Ethical Hacker (CEH)",
                        "CompTIA Security+",
                        "Offensive Security Certified Professional (OSCP)",
                        "Cloud Security Certifications",
                        "Security Operations and Incident Response"
                    ]
                else:
                    # Generic professional development courses
                    role_specific_courses = [
                        "LinkedIn Learning - Professional Skills Development",
                        "Coursera - Career Development Specialization",
                        "Udemy - Job Interview Skills Training",
                        "Project Management Professional (PMP)",
                        "Leadership and Management Skills",
                        "Technical Writing and Communication"
                    ]
                
                # Create a table for role-specific courses
                course_data = []
                for course in role_specific_courses:
                    course_data.append([Paragraph(f"• {clean_markdown(course)}", list_item_style)])
                
                course_table = Table(course_data, colWidths=[6*inch])
                course_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, 0), colors.lightblue),
                    ('TEXTCOLOR', (0, 0), (0, 0), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica'),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ]))
                
                content.append(course_table)
            
            content.append(Spacer(1, 0.2*inch))
            
            # Add footer with page numbers
            def add_page_number(canvas, doc):
                canvas.saveState()
                canvas.setFont('Helvetica', 9)
                page_num = canvas.getPageNumber()
                text = f"Page {page_num}"
                canvas.drawRightString(7.5*inch, 0.25*inch, text)
                
                # Add generation date at the bottom
                canvas.setFont('Helvetica', 9)
                date_text = f"Generated on: {datetime.datetime.now().strftime('%B %d, %Y')}"
                canvas.drawString(0.5*inch, 0.25*inch, date_text)
                
                canvas.restoreState()
            
            # Build the PDF
            doc.build(content, onFirstPage=add_page_number, onLaterPages=add_page_number)
            
            # Get the PDF from the buffer
            buffer.seek(0)
            return buffer
        
        except Exception as e:
            st.error(f"Error generating simple PDF report: {str(e)}")
            import traceback
            st.code(traceback.format_exc())
            return None 

    def process_sections(self, analysis_text, content, normal_style, list_item_style, subheading_style, heading_style, clean_markdown):
        """Process sections of the analysis text with special handling for certain sections"""
        # Parse the markdown-like content
        sections = analysis_text.split("##")
        
        # Define sections to include in detailed analysis
        detailed_sections = [
            "Professional Profile Analysis",
            "Skills Analysis",
            "Experience Analysis",
            "Education Analysis",
            "ATS Optimization Assessment",
            "Role Alignment Analysis",
            "Job Match Analysis"
        ]
        
        # Add Detailed Analysis section
        content.append(Paragraph("Detailed Analysis", heading_style))
        content.append(Spacer(1, 0.1*inch))
        
        for section in sections:
            if not section.strip():
                continue
            
            # Extract section title and content
            lines = section.strip().split("\n")
            section_title = lines[0].strip()
            
            # Skip sections we don't want in the detailed analysis
            if section_title not in detailed_sections and section_title != "Overall Assessment":
                continue
            
            # Skip Overall Assessment as we've already included it
            if section_title == "Overall Assessment":
                continue
            
            section_content = "\n".join(lines[1:]).strip()
            
            # Add section title
            content.append(Paragraph(section_title, subheading_style))
            content.append(Spacer(1, 0.1*inch))
            
            # Process content based on section
            if section_title == "Skills Analysis":
                # Extract current and missing skills
                current_skills = []
                missing_skills = []
                
                if "Current Skills" in section_content:
                    current_part = section_content.split("Current Skills")[1]
                    if "Missing Skills" in current_part:
                        current_part = current_part.split("Missing Skills")[0]
                    
                    for line in current_part.split("\n"):
                        if line.strip() and ("-" in line or "*" in line or "•" in line):
                            skill = clean_markdown(line.replace("-", "").replace("*", "").replace("•", "").strip())
                            if skill:
                                current_skills.append(skill)
                
                if "Missing Skills" in section_content:
                    missing_part = section_content.split("Missing Skills")[1]
                    for line in missing_part.split("\n"):
                        if line.strip() and ("-" in line or "*" in line or "•" in line):
                            skill = clean_markdown(line.replace("-", "").replace("*", "").replace("•", "").strip())
                            if skill:
                                missing_skills.append(skill)
                
                # Create skills table with better formatting
                if current_skills or missing_skills:
                    # Create paragraphs for each skill to ensure proper wrapping
                    current_skill_paragraphs = [Paragraph(skill, normal_style) for skill in current_skills]
                    missing_skill_paragraphs = [Paragraph(skill, normal_style) for skill in missing_skills]
                    
                    # Make sure both lists have the same length
                    max_len = max(len(current_skill_paragraphs), len(missing_skill_paragraphs))
                    current_skill_paragraphs.extend([Paragraph("", normal_style)] * (max_len - len(current_skill_paragraphs)))
                    missing_skill_paragraphs.extend([Paragraph("", normal_style)] * (max_len - len(missing_skill_paragraphs)))
                    
                    # Create data for the table
                    data = [["Current Skills", "Missing Skills"]]
                    for i in range(max_len):
                        data.append([current_skill_paragraphs[i], missing_skill_paragraphs[i]])
                    
                    # Create the table with fixed column widths
                    table = Table(data, colWidths=[3*inch, 3*inch])
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (1, 0), colors.lightgreen),
                        ('TEXTCOLOR', (0, 0), (1, 0), colors.black),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('LEFTPADDING', (0, 0), (-1, -1), 10),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                    ]))
                    
                    content.append(table)
                
                # We no longer need to add skill proficiency outside the table
                # as it's now included in the table itself
            elif section_title == "ATS Optimization Assessment":
                # Special handling for ATS Optimization Assessment
                ats_score_line = ""
                ats_content = []
                
                # Extract ATS score if present
                for line in section_content.split("\n"):
                    if "ATS Score:" in line:
                        ats_score_line = clean_markdown(line)
                    elif line.strip():
                        # Check if it's a list item
                        if line.strip().startswith("-") or line.strip().startswith("*") or line.strip().startswith("•"):
                            ats_content.append("• " + clean_markdown(line.strip()[1:].strip()))
                        else:
                            ats_content.append(clean_markdown(line))
                
                # Add ATS score line if found
                if ats_score_line:
                    content.append(Paragraph(ats_score_line, normal_style))
                    content.append(Spacer(1, 0.1*inch))
                
                # Add the rest of the ATS content
                for para in ats_content:
                    if para.startswith("• "):
                        content.append(Paragraph(para, list_item_style))
                    else:
                        content.append(Paragraph(para, normal_style))
            else:
                # Process regular paragraphs
                paragraphs = section_content.split("\n")
                for para in paragraphs:
                    if para.strip():
                        # Check if it's a list item
                        if para.strip().startswith("-") or para.strip().startswith("*") or para.strip().startswith("•"):
                            para = "• " + clean_markdown(para.strip()[1:].strip())
                            content.append(Paragraph(para, list_item_style))
                        else:
                            content.append(Paragraph(clean_markdown(para), normal_style))
            
            content.append(Spacer(1, 0.2*inch))
        
        return content
